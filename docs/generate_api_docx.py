from pathlib import Path

from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parent
SOURCE = ROOT / "API_DOCUMENTATION.md"
TARGET = ROOT / "API_DOCUMENTATION.docx"


def add_markdown_like_content(doc: Document, line: str) -> None:
    stripped = line.rstrip()
    if not stripped:
        doc.add_paragraph("")
        return

    if stripped.startswith("### "):
        doc.add_heading(stripped[4:], level=3)
        return
    if stripped.startswith("## "):
        doc.add_heading(stripped[3:], level=2)
        return
    if stripped.startswith("# "):
        doc.add_heading(stripped[2:], level=1)
        return
    if stripped.startswith("- "):
        doc.add_paragraph(stripped[2:], style="List Bullet")
        return

    doc.add_paragraph(stripped)


def main() -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)

    if not SOURCE.exists():
        raise FileNotFoundError(f"Missing source markdown: {SOURCE}")

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        add_markdown_like_content(doc, raw_line)

    doc.save(TARGET)
    print(f"Generated: {TARGET}")


if __name__ == "__main__":
    main()
