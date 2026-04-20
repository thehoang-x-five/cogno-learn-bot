# -*- coding: utf-8 -*-
"""
Step 1: Fix corrupted docx (NULL image reference)
Step 2: Apply all review fixes to the report
"""
import sys, os, shutil, zipfile, re, copy
from pathlib import Path

sys.stdout.reconfigure(encoding="utf-8")
sys.stderr.reconfigure(encoding="utf-8")

SRC = Path(r"c:\Users\THINKPAD\Documents\GitHub\chatbot\finalfullstack1.docx")
BAK = SRC.with_suffix(".docx.bak")
FIXED = SRC.with_name("finalfullstack1_fixed.docx")

# ── Step 1: Fix corrupted NULL reference in rels ──────────────
print("=" * 60)
print("STEP 1: Fixing corrupted NULL image reference...")
print("=" * 60)

# Back up original
if not BAK.exists():
    shutil.copy2(SRC, BAK)
    print(f"  ✅ Backup created: {BAK.name}")

# Read and fix the rels XML
with zipfile.ZipFile(SRC, "r") as zin:
    entries = {}
    for name in zin.namelist():
        entries[name] = zin.read(name)

rels_key = "word/_rels/document.xml.rels"
rels_xml = entries[rels_key].decode("utf-8")

# Remove the broken relationship pointing to ../NULL
rels_xml_fixed = re.sub(
    r'<Relationship[^>]*Target="\.\./NULL"[^/]*/>', "", rels_xml
)

if rels_xml != rels_xml_fixed:
    print("  ✅ Removed broken NULL relationship")
else:
    print("  ℹ️  No NULL reference found (already fixed?)")

entries[rels_key] = rels_xml_fixed.encode("utf-8")

# Write fixed zip
with zipfile.ZipFile(FIXED, "w", zipfile.ZIP_DEFLATED) as zout:
    for name, data in entries.items():
        zout.writestr(name, data)
print(f"  ✅ Fixed file written: {FIXED.name}")

# ── Step 2: Apply all review fixes ────────────────────────────
print()
print("=" * 60)
print("STEP 2: Applying report fixes...")
print("=" * 60)

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(str(FIXED))

changes_made = []

# Helper: find & replace in all paragraphs
def find_replace_text(old, new, limit=0):
    """Find and replace text across all paragraphs. Returns count of replacements."""
    count = 0
    for para in doc.paragraphs:
        if old in para.text:
            for run in para.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)
                    count += 1
                    if limit and count >= limit:
                        return count
    # Also search in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old in para.text:
                        for run in para.runs:
                            if old in run.text:
                                run.text = run.text.replace(old, new)
                                count += 1
    return count

def find_replace_in_runs(paragraph, old_text, new_text):
    """More robust find-replace that handles text split across runs."""
    full = paragraph.text
    if old_text not in full:
        return False
    # Try run-level first
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    # If text spans runs, rebuild
    new_full = full.replace(old_text, new_text)
    if new_full != full:
        # Clear all runs except first, put all text in first run
        if paragraph.runs:
            first_run = paragraph.runs[0]
            fmt = first_run.font
            for run in paragraph.runs[1:]:
                run.text = ""
            first_run.text = new_full
            return True
    return False


# ── Fix 1: Title spacing ──
print("\n[Fix 1] Fixing title spacing (AI**HỖ** → AI HỖ)...")
c = find_replace_text("AI\u200bHỖ", "AI HỖ")  # zero-width space
c += find_replace_text("AIHỖ", "AI HỖ")
c += find_replace_text("AI HỖ", "AI HỖ")  # double space fix
for para in doc.paragraphs:
    txt = para.text
    if "TÍCH HỢP AI" in txt and "HỖ TRỢ" in txt:
        # Check if AI and HỖ are in different runs but no space between them
        runs_text = "".join(r.text for r in para.runs)
        if "AIHỖ" in runs_text or "AI HỖ" not in runs_text:
            # Fix across runs
            for i, run in enumerate(para.runs):
                if run.text.endswith("AI") and i + 1 < len(para.runs):
                    next_run = para.runs[i + 1]
                    if next_run.text.startswith("HỖ"):
                        run.text = run.text + " "
                        c += 1
                        break
if c:
    changes_made.append(f"Fix 1: Sửa tiêu đề thiếu dấu cách ({c} chỗ)")
    print(f"  ✅ Fixed {c} occurrences")
else:
    print("  ℹ️  Not found or already fixed")

# ── Fix 2: Fix GVHD name ──
print("\n[Fix 2] Fixing teacher name (Nguyễn Thị Hường → ThS. Vũ Đình Hồng)...")
c = find_replace_text("Nguyễn Thị Hường", "ThS. Vũ Đình Hồng")
c += find_replace_text("cô Nguyễn Thị Hường", "thầy Vũ Đình Hồng")
c += find_replace_text("cô Vũ Đình Hồng", "thầy Vũ Đình Hồng")  # in case
if c:
    changes_made.append(f"Fix 2: Sửa tên GVHD ({c} chỗ)")
    print(f"  ✅ Fixed {c} occurrences")
else:
    print("  ℹ️  Not found or already fixed")

# ── Fix 3: Fix Next.js → Vite ──
print("\n[Fix 3] Fixing Next.js → Vite...")
c = 0
c += find_replace_text("Next.js/React", "Vite + React")
c += find_replace_text("Next.js (Vite), React", "Vite, React")
c += find_replace_text("Next.js (Vite)", "Vite")
c += find_replace_text("Next.js, React", "Vite, React")
# Be careful: don't replace "Next.js" in references/bibliography
for para in doc.paragraphs:
    if "Next.js" in para.text:
        # Skip if it's in a reference/bibliography context
        context = para.text.lower()
        if "tài liệu tham khảo" in context or "[" in para.text[:5]:
            continue
        for run in para.runs:
            if "Next.js" in run.text:
                run.text = run.text.replace("Next.js", "Vite")
                c += 1
if c:
    changes_made.append(f"Fix 3: Sửa Next.js → Vite ({c} chỗ)")
    print(f"  ✅ Fixed {c} occurrences")
else:
    print("  ℹ️  Not found or already fixed")

# ── Fix 4: Fix member names stuck together ──
print("\n[Fix 4] Fixing member names stuck together...")
c = 0
stuck_patterns = [
    ("TRƯƠNG DƯƠNG HƯNGHOÀNG CÔNG TÀI THẾ", "TRƯƠNG DƯƠNG HƯNG\nHOÀNG CÔNG TÀI THẾ"),
    ("HOÀNG CÔNG TÀI THẾNGUYỄN THẾ AN", "HOÀNG CÔNG TÀI THẾ\nNGUYỄN THẾ AN"),
    ("TRƯƠNG DƯƠNG HƯNG HOÀNG CÔNG TÀI THẾ", "TRƯƠNG DƯƠNG HƯNG\nHOÀNG CÔNG TÀI THẾ"),
]
for old, new in stuck_patterns:
    c += find_replace_text(old, new)
if c:
    changes_made.append(f"Fix 4: Tách tên thành viên bị dính ({c} chỗ)")
    print(f"  ✅ Fixed {c} occurrences")
else:
    print("  ℹ️  Not found or already fixed")

# ── Fix 5: Fix "state client" → better phrasing ──
print("\n[Fix 5] Fixing 'state client' wording...")
c = find_replace_text("state client", "trạng thái ứng dụng (React state, context)")
if c:
    changes_made.append(f"Fix 5: Sửa 'state client' ({c} chỗ)")
    print(f"  ✅ Fixed {c} occurrences")
else:
    print("  ℹ️  Not found or already fixed")

# ── Fix 6: Fix "Node 20+" → "Node.js 20+" ──
print("\n[Fix 6] Fixing 'Node 20+' → 'Node.js 20+'...")
c = find_replace_text("Node 20+", "Node.js phiên bản 20 trở lên")
c += find_replace_text("Node 20", "Node.js 20")
if c:
    changes_made.append(f"Fix 6: Sửa 'Node 20+' ({c} chỗ)")
    print(f"  ✅ Fixed {c} occurrences")
else:
    print("  ℹ️  Not found or already fixed")

# ── Fix 7: Rewrite LỜI CẢM ƠN ──
print("\n[Fix 7] Rewriting LỜI CẢM ƠN...")
cam_on_new = [
    "Em xin gửi lời cảm ơn chân thành đến toàn thể quý thầy cô Khoa Công nghệ Thông tin, Trường Đại học Tôn Đức Thắng đã tạo điều kiện thuận lợi cho em trong suốt quá trình học tập và thực hiện đồ án này.",
    "",
    "Đặc biệt, em xin bày tỏ lòng biết ơn sâu sắc đến ThS. Vũ Đình Hồng — giảng viên hướng dẫn học phần Phát triển ứng dụng Full-stack. Thầy không chỉ truyền đạt kiến thức nền tảng về kiến trúc phần mềm, tích hợp API và triển khai container mà còn định hướng cho nhóm em tiếp cận các kỹ thuật hiện đại như Retrieval-Augmented Generation (RAG), streaming SSE và hybrid search — những nội dung vượt ra ngoài khuôn khổ giáo trình truyền thống nhưng lại rất thiết thực cho nghề nghiệp tương lai.",
    "",
    "Nhờ những kiến thức từ học phần Full-stack, em đã có cơ hội áp dụng đồng thời React, FastAPI, Docker Compose cùng nhiều dịch vụ dữ liệu (MySQL, ChromaDB, Redis) vào một hệ thống hoàn chỉnh, qua đó hiểu sâu hơn về cách các thành phần frontend, backend, cơ sở dữ liệu và pipeline AI phối hợp trong thực tế.",
    "",
    "Trong quá trình thực hiện, do kiến thức và kinh nghiệm thực tiễn còn hạn chế nên kết quả không tránh khỏi thiếu sót. Em kính mong nhận được ý kiến đóng góp của thầy cô để hoàn thiện kỹ năng và rút kinh nghiệm cho các dự án sau.",
    "",
    "Kính chúc quý thầy cô dồi dào sức khỏe, đạt nhiều thành tựu trong công tác giảng dạy và nghiên cứu!",
]

cam_on_found = False
for i, para in enumerate(doc.paragraphs):
    txt = para.text.strip().upper()
    if "LỜI CẢM ƠN" in txt and len(txt) < 30:
        cam_on_found = True
        # Find the content paragraphs after the title
        j = i + 1
        # Clear existing content until next major heading
        content_cleared = 0
        while j < len(doc.paragraphs):
            next_para = doc.paragraphs[j]
            next_txt = next_para.text.strip().upper()
            # Stop at next chapter/section header
            if next_txt and (
                next_txt.startswith("CHƯƠNG") or
                next_txt.startswith("MỤC LỤC") or
                next_txt.startswith("LỜI CAM ĐOAN") or
                "NHẬN XÉT" in next_txt or
                next_txt.startswith("TÓM TẮT") or
                next_txt.startswith("DANH MỤC")
            ):
                break
            if next_para.text.strip():
                # Clear existing content
                for run in next_para.runs:
                    run.text = ""
                content_cleared += 1
            j += 1
        
        # Now write new content into the cleared paragraphs
        insert_idx = i + 1
        for k, line in enumerate(cam_on_new):
            if insert_idx + k < len(doc.paragraphs):
                target = doc.paragraphs[insert_idx + k]
                if target.runs:
                    target.runs[0].text = line
                    # Clear other runs 
                    for r in target.runs[1:]:
                        r.text = ""
                else:
                    target.add_run(line)
        
        changes_made.append("Fix 7: Viết lại LỜI CẢM ƠN")
        print(f"  ✅ Rewrote LỜI CẢM ƠN (cleared {content_cleared} paragraphs)")
        break

if not cam_on_found:
    print("  ⚠️  LỜI CẢM ƠN section not found")

# ── Fix 8: Rewrite CHƯƠNG 6 — KẾT LUẬN ──
print("\n[Fix 8] Rewriting Chương 6 — Kết luận...")
ket_luan_content = {
    "6.1": [
        "Qua quá trình thực hiện đồ án, nhóm em đã xây dựng thành công hệ thống web trợ lý học tập tích hợp AI, đáp ứng đầy đủ bốn cấp độ yêu cầu của đề bài môn Phát triển ứng dụng Full-stack.",
        "",
        "Ở tầng nền tảng, hệ thống cung cấp xác thực Google OAuth kết hợp JWT và phân quyền ba vai trò (Admin, Giáo viên, Sinh viên), quản lý khóa học với import Excel hàng loạt, cùng giao diện React/Vite hiện đại và responsive.",
        "",
        "Ở tầng RAG, nhóm em đã hiện thực pipeline truy xuất lai (hybrid search) kết hợp vector similarity và BM25, hợp nhất bằng Reciprocal Rank Fusion, tinh chỉnh qua cross-encoder reranking, và sinh câu trả lời streaming kèm trích dẫn nguồn tài liệu — giúp sinh viên kiểm chứng thông tin trực tiếp thay vì phụ thuộc hoàn toàn vào mô hình ngôn ngữ.",
        "",
        "Ở tầng trải nghiệm, chat hỗ trợ SSE streaming cho phản hồi dần, bộ nhớ cuộn (rolling summary) duy trì ngữ cảnh hội thoại dài, và cấu hình model linh hoạt qua giao diện quản trị.",
        "",
        "Ở tầng nâng cao, agent orchestrator với Gemini function calling cho phép chatbot tự quyết định gọi tool tra cứu lịch thi từ cơ sở dữ liệu, quiz AI tự sinh đề từ nội dung tài liệu đã upload, và ingestion bất đồng bộ qua Celery xử lý tài liệu nặng mà không chặn API.",
        "",
        "Toàn bộ hệ thống được đóng gói Docker Compose (MySQL, ChromaDB, Redis, backend, worker, frontend) và có thể demo end-to-end trên một máy duy nhất chỉ với một lệnh docker compose up.",
    ],
    "6.2": [
        "Chất lượng câu trả lời RAG phụ thuộc trực tiếp vào corpus đầu vào: tài liệu PDF scan chất lượng thấp, slide chứa nhiều hình ảnh mà ít text, hoặc cấu hình chunk/embedding chưa tối ưu đều ảnh hưởng đến độ chính xác của truy xuất. Với các truy vấn mơ hồ, hệ thống vẫn có thể trả kết quả không ổn định.",
        "",
        "Chi phí và giới hạn tốc độ của Gemini API ảnh hưởng trải nghiệm khi nhiều người dùng đồng thời hoặc prompt quá dài. Kiểm thử tự động mới bao phủ một số API chính, chưa có E2E test giao diện và chưa mô phỏng được kịch bản đa người dùng tải cao.",
        "",
        "Khía cạnh bảo mật cho triển khai thực tế (rate limiting, WAF, xoay secret định kỳ) mới dừng ở mức nhận diện trong báo cáo, chưa được triển khai đầy đủ trong phạm vi đồ án.",
    ],
    "6.3": [
        "Về RAG và đánh giá: xây dựng bộ benchmark offline với các metric như hit@k, MRR để đo lường chất lượng truy xuất một cách định lượng; bổ sung logging trace cho từng bước retrieval/rerank nhằm phát hiện điểm yếu; và thử nghiệm A/B giữa các chiến lược prompt khác nhau.",
        "",
        "Về sản phẩm: tích hợp E2E testing bằng Playwright để tự động hóa kiểm thử giao diện; hỗ trợ đa ngôn ngữ (i18n); tối ưu bundle frontend và giám sát lỗi runtime qua Sentry hoặc OpenTelemetry.",
        "",
        "Về hạ tầng: thiết lập pipeline CI/CD, sao lưu MySQL và ChromaDB định kỳ, mở rộng ngang (horizontal scale) cho Celery worker khi lượng tài liệu tăng.",
        "",
        "Về nghiệp vụ giáo dục: phát triển module ôn tập cá nhân hóa (adaptive review), tích hợp với hệ thống LMS hiện có, và phân tích hành vi học tập của sinh viên (tuân thủ quy định bảo vệ dữ liệu cá nhân).",
    ],
}

# Find chapter 6 sections and rewrite
for section_key, new_lines in ket_luan_content.items():
    for i, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if txt.startswith(section_key) and (
            "Kết quả" in txt or "Hạn chế" in txt or "Hướng phát triển" in txt or
            "kết quả" in txt or "hạn chế" in txt or "hướng" in txt
        ):
            # Found section header, rewrite content after it
            j = i + 1
            line_idx = 0
            while j < len(doc.paragraphs) and line_idx < len(new_lines):
                next_para = doc.paragraphs[j]
                next_txt = next_para.text.strip()
                # Stop at next section header
                if next_txt and (
                    re.match(r"^6\.\d", next_txt) or
                    next_txt.upper().startswith("CHƯƠNG") or
                    next_txt.upper().startswith("TÀI LIỆU")
                ):
                    break
                # Rewrite this paragraph
                if next_para.runs:
                    next_para.runs[0].text = new_lines[line_idx]
                    for r in next_para.runs[1:]:
                        r.text = ""
                else:
                    next_para.add_run(new_lines[line_idx])
                line_idx += 1
                j += 1
            if line_idx > 0:
                print(f"  ✅ Rewrote section {section_key} ({line_idx} lines)")
            break

changes_made.append("Fix 8: Viết lại Chương 6 — Kết luận")

# ── Fix 9: Add phân công công việc table reference ──
print("\n[Fix 9] Checking phân công công việc...")
has_phancong = False
for para in doc.paragraphs:
    if "phân công" in para.text.lower() and ("công việc" in para.text.lower() or "thành viên" in para.text.lower()):
        has_phancong = True
        break
if not has_phancong:
    print("  ⚠️  Bảng phân công chưa có — cần thêm thủ công trong Word")
    changes_made.append("Fix 9: [CẦN THỦ CÔNG] Thêm bảng phân công công việc")
else:
    print("  ℹ️  Already has phân công section")

# ── Fix 10: Fix test file references ──
print("\n[Fix 10] Fixing test file references...")
c = 0
test_fixes = [
    ("test_pdf_citation_pipeline.py", "test_rag_pipeline.py"),
    ("test_agent_function_calling.py", "test_agent_orchestrator.py"),
]
for old, new in test_fixes:
    c += find_replace_text(old, new)
if c:
    changes_made.append(f"Fix 10: Sửa tên file test ({c} chỗ)")
    print(f"  ✅ Fixed {c} test file references")
else:
    print("  ℹ️  Not found or already fixed")

# ── Save ──
print()
print("=" * 60)
print("SAVING...")
print("=" * 60)

output_path = SRC.with_name("finalfullstack1_v2.docx")
doc.save(str(output_path))
print(f"\n✅ Saved to: {output_path.name}")
print(f"📋 Original backed up to: {BAK.name}")

print()
print("=" * 60)
print("SUMMARY OF CHANGES")
print("=" * 60)
for ch in changes_made:
    print(f"  ✔ {ch}")

print()
print("⚠️  CÁC MỤC CẦN CHỈNH THỦ CÔNG TRONG WORD:")
print("  1. Xóa trang bìa bị lặp (nếu có)")
print("  2. Xóa bảng API trùng lặp (Bảng 3.2 vs 3.4)")
print("  3. Thêm DANH MỤC HÌNH ẢNH, DANH MỤC BẢNG BIỂU sau Mục lục")
print("  4. Thêm bảng phân công công việc")
print("  5. Thêm cột 'Kết quả thực tế' và 'Đạt' vào bảng kiểm thử 5.3")
print("  6. Đánh số tài liệu tham khảo [1], [2]... theo IEEE")
print("  7. Chèn ảnh sơ đồ từ docs/diagrams (render PlantUML → PNG)")
print("  8. Kiểm tra lại format, font chữ sau khi mở file v2")
